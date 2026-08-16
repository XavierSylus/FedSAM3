from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path(__file__).resolve().parents[1]
USABLE_WIDTH_DXA = 9360


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Build an editable ICASSP-style Word manuscript from the verified LaTeX source."
    )
    parser.add_argument("--config", type=Path, required=True)
    return parser.parse_args()


def read_config(config_path: Path) -> dict:
    return json.loads(config_path.read_text(encoding="utf-8"))


def resolve_workspace_path(value: str) -> Path:
    return WORKSPACE_ROOT / Path(value)


def set_run_font(run, name: str, size: float, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top: int = 55, start: int = 75, bottom: int = 55, end: int = 75) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != USABLE_WIDTH_DXA:
        raise ValueError("Table widths must add up to the usable page width.")
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    table_layout = table_pr.first_child_found_in("w:tblLayout")
    if table_layout is None:
        table_layout = OxmlElement("w:tblLayout")
        table_pr.append(table_layout)
    table_layout.set(qn("w:type"), "fixed")

    table_grid = table._tbl.tblGrid
    for grid_col, width in zip(table_grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.tcW
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_section_geometry(section, config: dict) -> None:
    margins = config["margins_inches"]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(margins["top"])
    section.right_margin = Inches(margins["right"])
    section.bottom_margin = Inches(margins["bottom"])
    section.left_margin = Inches(margins["left"])
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)


def set_columns(section, count: int, gap_inches: float) -> None:
    section_pr = section._sectPr
    columns = section_pr.first_child_found_in("w:cols")
    if columns is None:
        columns = OxmlElement("w:cols")
        section_pr.append(columns)
    columns.set(qn("w:num"), str(count))
    columns.set(qn("w:space"), str(round(gap_inches * 1440)))
    columns.set(qn("w:equalWidth"), "1")


def add_section_break(doc: Document, config: dict, columns: int) -> None:
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_geometry(section, config)
    set_columns(section, columns, config["column_gap_inches"])


def configure_styles(doc: Document, config: dict) -> None:
    font = config["body_font"]
    body_size = config["body_size_pt"]
    normal = doc.styles["Normal"]
    normal.font.name = font
    normal._element.rPr.rFonts.set(qn("w:ascii"), font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    normal.font.size = Pt(body_size)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, size in (("Heading 1", 9.5), ("Heading 2", 9.2)):
        style = doc.styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:ascii"), font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(6 if style_name == "Heading 1" else 4)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_with_next = True


def add_paragraph(doc: Document, text: str, config: dict, *, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after: float = 2, size: float | None = None, italic: bool = False, bold: bool = False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, config["body_font"], size or config["body_size_pt"], bold=bold, italic=italic)
    return paragraph


def add_heading(doc: Document, text: str, config: dict, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, config["body_font"], 9.5 if level == 1 else 9.2, bold=True)
    if level == 1:
        run.text = text.upper()
    return paragraph


def add_equation(doc: Document, text: str, number: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(f"{text}    ({number})")
    set_run_font(run, "Cambria Math", 8.5)


def clean_citation_text(text: str, citation_numbers: dict[str, int]) -> str:
    def replace_citation(match: re.Match[str]) -> str:
        keys = [key.strip() for key in match.group(1).split(",")]
        return "[" + ", ".join(str(citation_numbers[key]) for key in keys) + "]"

    text = re.sub(r"\\cite\{([^}]+)\}", replace_citation, text)
    text = text.replace("$", "")
    text = text.replace("\\times", "×").replace("\\circ", "°")
    text = text.replace("\\lambda_{\\mathrm{cream}}", "λcream")
    text = text.replace("\\mu", "μ").replace("^{-5}", "⁻⁵")
    text = text.replace("B--C", "B-C")
    return text


def extract_tex_metadata(tex: str) -> tuple[str, str, list[str]]:
    title = re.search(r"\\title\{(.+)\}", tex).group(1)
    author = re.search(r"\\name\{(.+)\}", tex).group(1)
    address_block = re.search(r"\\address\{(.+?)\}", tex, flags=re.DOTALL).group(1)
    address_lines = [line.strip() for line in address_block.split("\\\\")]
    return title, author, address_lines


def parse_bibliography(bib_text: str) -> dict[str, dict[str, str]]:
    entries: dict[str, dict[str, str]] = {}
    start_pattern = re.compile(r"@(\w+)\{([^,]+),")
    cursor = 0
    while match := start_pattern.search(bib_text, cursor):
        entry_type, key = match.groups()
        index = match.end()
        depth = 1
        while index < len(bib_text) and depth:
            if bib_text[index] == "{":
                depth += 1
            elif bib_text[index] == "}":
                depth -= 1
            index += 1
        body = bib_text[match.end() : index - 1]
        fields = {"entry_type": entry_type}
        for field_match in re.finditer(r"(\w+)\s*=\s*\{([^{}]*)\}", body, flags=re.DOTALL):
            fields[field_match.group(1).lower()] = re.sub(r"\s+", " ", field_match.group(2)).strip()
        entries[key.strip()] = fields
        cursor = index
    return entries


def citation_order(tex: str) -> list[str]:
    ordered: list[str] = []
    for group in re.findall(r"\\cite\{([^}]+)\}", tex):
        for key in (item.strip() for item in group.split(",")):
            if key not in ordered:
                ordered.append(key)
    return ordered


def simplify_bib_text(text: str) -> str:
    replacements = {
        "\\&": "&",
        "{": "",
        "}": "",
        "\\textit": "",
        "\\textbf": "",
        "--": "-",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return re.sub(r"\s+", " ", text).strip()


def format_reference(index: int, entry: dict[str, str]) -> str:
    authors = simplify_bib_text(entry.get("author", ""))
    title = simplify_bib_text(entry.get("title", ""))
    venue = simplify_bib_text(entry.get("journal") or entry.get("booktitle", ""))
    parts = [f"[{index}] {authors}", f'“{title},"']
    if venue:
        parts.append(venue)
    if entry.get("volume"):
        parts.append(f"vol. {entry['volume']}")
    if entry.get("number"):
        parts.append(f"no. {entry['number']}")
    if entry.get("pages"):
        parts.append(f"pp. {simplify_bib_text(entry['pages'])}")
    if entry.get("year"):
        parts.append(entry["year"])
    return ", ".join(part for part in parts if part) + "."


def write_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], config: dict, caption: str, subheads: dict[int, str] | None = None) -> None:
    caption_paragraph = add_paragraph(doc, caption, config, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=2, size=8, bold=True)
    caption_paragraph.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, headers):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, config["body_font"], 7.7, bold=True)
    set_repeat_table_header(table.rows[0])

    for row_index, row_values in enumerate(rows):
        if subheads and row_index in subheads:
            row = table.add_row()
            merged = row.cells[0]
            for cell in row.cells[1:]:
                merged = merged.merge(cell)
            paragraph = merged.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(subheads[row_index])
            set_run_font(run, config["body_font"], 7.6, bold=True)
        row = table.add_row()
        for index, (cell, text) in enumerate(zip(row.cells, row_values)):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index < 2 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, config["body_font"], 7.5)
    add_paragraph(doc, "", config, after=1, size=1)


def add_framework_figure(doc: Document, image_path: Path, config: dict) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)
    shape = paragraph.add_run().add_picture(str(image_path), width=Inches(3.1))
    shape._inline.docPr.set("descr", "FedSAM3-Hetero framework overview")
    add_paragraph(
        doc,
        "Figure 1. FedSAM3-Hetero overview. Restricted routing is active only in Group C; server representations are auxiliary and do not determine routing.",
        config,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
        size=8,
    )


def build_document(config: dict) -> Path:
    source_tex = resolve_workspace_path(config["source_tex"])
    source_bib = resolve_workspace_path(config["source_bib"])
    framework_image = resolve_workspace_path(config["framework_image"])
    output_docx = resolve_workspace_path(config["output_docx"])
    tex = source_tex.read_text(encoding="utf-8")
    bibliography = parse_bibliography(source_bib.read_text(encoding="utf-8"))
    order = citation_order(tex)
    citation_numbers = {key: index for index, key in enumerate(order, start=1)}
    title, author, address_lines = extract_tex_metadata(tex)

    doc = Document()
    set_section_geometry(doc.sections[0], config)
    set_columns(doc.sections[0], 1, config["column_gap_inches"])
    configure_styles(doc, config)
    properties = doc.core_properties
    properties.title = title
    properties.author = author
    properties.subject = "Editable ICASSP-style manuscript"

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(0)
    title_paragraph.paragraph_format.space_after = Pt(6)
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(title.upper())
    set_run_font(title_run, config["body_font"], 14, bold=True)

    author_paragraph = doc.add_paragraph()
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_paragraph.paragraph_format.space_after = Pt(3)
    author_run = author_paragraph.add_run(author)
    set_run_font(author_run, config["body_font"], 10, italic=True)
    for line in address_lines:
        address_paragraph = doc.add_paragraph()
        address_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        address_paragraph.paragraph_format.space_after = Pt(0)
        address_run = address_paragraph.add_run(line)
        set_run_font(address_run, config["body_font"], 8.5)

    add_section_break(doc, config, columns=2)
    add_heading(doc, "Abstract", config)
    add_paragraph(
        doc,
        "Federated brain tumor segmentation must accommodate institutions with unequal imaging modalities, annotations, and local objectives. Existing aggregation often assumes a common update space and can directly mix updates whose parameter responsibilities differ. We present FedSAM3-Hetero, an experimental framework for missing-modality heterogeneous collaboration. The framework partitions parameters by semantic responsibility and applies server-side restricted routing: a client contributes only to groups that its supervision can reliably update. A server image/text representation path provides an auxiliary cross-round semantic reference. On a BraTS-style MRI task, multimodal participation improves the image-only FedAvg baseline in Dice and HD95. With the same client composition, restricted routing leaves Dice essentially unchanged while providing small, directionally consistent boundary-quality and training-stability benefits. An external FedProx setting obtains stronger endpoint metrics but changes the local objective, so it is not a same-loss ablation. The framework therefore supports analysis of parameter-mixing boundaries rather than claiming universal superiority over robust federated optimizers.",
        config,
        after=2,
    )
    keywords = doc.add_paragraph()
    keywords.paragraph_format.space_after = Pt(3)
    label = keywords.add_run("Index Terms—")
    set_run_font(label, config["body_font"], 9, italic=True, bold=True)
    terms = keywords.add_run("Federated learning, medical image segmentation, missing modality, heterogeneous clients, parameter routing, SAM3")
    set_run_font(terms, config["body_font"], 9, italic=True)

    add_heading(doc, "1. Introduction", config)
    add_paragraph(doc, clean_citation_text("Brain tumor segmentation supports lesion localization, volume estimation, and treatment planning. Modern U-Net, Transformer, and foundation-model approaches improve anatomical and pathological representation under centralized access \\cite{nnunet,transunet,sam2,sammed2d,sam3}. In clinical collaboration, however, privacy constraints keep data at separate institutions. Federated learning (FL) enables joint training without moving those data \\cite{fedavg,kairouz}.", citation_numbers), config)
    add_paragraph(doc, "Most federated medical-segmentation methods assume that clients share compatible modalities, supervision, objectives, and trainable parameter spaces. This assumption fails when institutions have different modality availability and local supervision. For example, image-only clients optimize visual segmentation, whereas image-text clients also optimize cross-modal consistency. Direct FedAvg over every parameter can then mix updates with different semantic responsibilities.", config)
    add_paragraph(doc, "This paper studies whether those updates should be aggregated indiscriminately. FedSAM3-Hetero defines client roles and parameter groups, then enforces server-side admission by group. The main experiments compare an image-only baseline, direct heterogeneous FedAvg, and the same heterogeneous composition with restricted routing. The contributions are:", config)
    contributions = [
        "A missing-modality heterogeneous FL framework that makes client roles, local objectives, and server aggregation explicit.",
        "Parameter-group restricted routing that excludes unreliable client updates from each group.",
        "Protocol-controlled evidence that distinguishes multimodal participation, routing, global representation update, and an external FedProx comparison.",
    ]
    for item in contributions:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(item)
        set_run_font(run, config["body_font"], 9)

    add_heading(doc, "2. Related Work", config)
    add_paragraph(doc, clean_citation_text("Federated medical segmentation addresses privacy-preserving multi-center learning under non-IID data and sample imbalance. FedProx, SCAFFOLD, and FedBN respectively target client drift, control variates, and feature-shifted normalization \\cite{fedprox,scaffold,fedbn}. Foundation-model transfer has also been explored for federated medical segmentation \\cite{fedfms}.", citation_numbers), config)
    add_paragraph(doc, clean_citation_text("Missing-modality collaboration is more specific than conventional statistical heterogeneity. FedAMM considers arbitrary missing modalities in federated brain tumor segmentation, while CAR-MFL uses retrieval-based cross-modal augmentation \\cite{fedamm,carmfl}. Recent ICASSP work illustrates complementary challenges: FedSLAG studies heterogeneous supervision \\cite{fedsag}; MM-DC learns a modality-invariant space for multi-modality semi-supervised segmentation \\cite{mmdc}; and SAM2-SP adapts a promptable foundation model to 3D medical segmentation \\cite{sam2sp}. In contrast, FedSAM3-Hetero focuses on whether a client's update is admissible for a parameter group under its local supervision.", citation_numbers), config)

    add_heading(doc, "3. FedSAM3-Hetero", config)
    add_heading(doc, "3.1. Problem formulation", config, level=2)
    add_paragraph(doc, "Let C = {c₁, c₂, …, cN} be the client set. The quantitative protocols contain image-only and multimodal clients, denoted by", config)
    add_equation(doc, "Cimg = {ci ∈ C | mi = img};  Cmm = {ci ∈ C | mi = mm};  Cimg ∪ Cmm ⊆ C;  Cimg ∩ Cmm = ∅.", 1)
    add_paragraph(doc, "At communication round t, the global parameters are partitioned as", config)
    add_equation(doc, "θᵗ = {θsharedᵗ, θimgᵗ, θmmᵗ}.", 2)
    add_paragraph(doc, "The subsets respectively represent shared, vision-driven, and multimodal semantic-interaction parameters. Image-only clients use segmentation loss; multimodal clients additionally use CREAM consistency and image-text alignment:", config)
    add_equation(doc, "Limg = Lseg;  Lmm = Lseg + λcream Lcream + λalign Lalign.", 3)

    add_heading(doc, "3.2. Restricted routing", config, level=2)
    add_paragraph(doc, "Standard FedAvg updates all parameters as", config)
    add_equation(doc, "θᵗ⁺¹ = Σₖ₌₁ᴺ wk θkᵗ⁺¹;  wk ≥ 0;  Σₖ₌₁ᴺ wk = 1.", 4)
    add_paragraph(doc, "FedSAM3-Hetero instead defines an eligible client set Sg for each parameter group g and aggregates only admitted uploads:", config)
    add_equation(doc, "θgᵗ⁺¹ = Σk∈Sg αk,g θk,gᵗ⁺¹;  αk,g ≥ 0;  Σk∈Sg αk,g = 1.", 5)
    add_paragraph(doc, "When no client is eligible, the server retains the previous value,", config)
    add_equation(doc, "θgᵗ⁺¹ = θgᵗ.", 6)
    add_paragraph(doc, "Thus, shared and visual groups accept visually supervised clients, while multimodal interaction groups accept only multimodal clients. The rule is server-side whitelist filtering and does not alter the SAM3 encoder or its geometric encoding.", config)

    add_heading(doc, "3.3. Auxiliary global representations", config, level=2)
    add_paragraph(doc, "The server also maintains image and text semantic references using EMA-normalized proxy representations:", config)
    add_equation(doc, "rimgᵗ⁺¹ = N(μ rimgᵗ + (1 − μ) r̄imgᵗ⁺¹);  rtxtᵗ⁺¹ = N(μ rtxtᵗ + (1 − μ) r̄txtᵗ⁺¹);  μ ∈ [0, 1].", 7)
    add_paragraph(doc, "These vectors do not decide routing; they are evaluated separately as an auxiliary component. Figure 1 summarizes the data and parameter flow before the experimental evidence below.", config)
    add_framework_figure(doc, framework_image, config)

    add_heading(doc, "4. Experiments", config)
    add_heading(doc, "4.1. Protocol", config, level=2)
    add_paragraph(doc, clean_citation_text("We evaluate BraTS-style MRI brain tumor segmentation \\cite{brats,tcga}. Textual semantics follow the motivation of TextBraTS \\cite{textbrats}; text-only clients are not included in the quantitative comparisons. All groups use input size 256, three classes, batch size one, four accumulation steps, effective batch size four, learning rate 5×10⁻⁵, and at most 60 communication rounds. Because the current labels are closer to BG, WT, and ET than to a complete WT/TC/ET loop, evaluation reports WT and ET behavior through Dice and HD95.", citation_numbers), config)
    add_paragraph(doc, "Group A is image-only FedAvg. Group B adds multimodal clients and still uses direct FedAvg. Group C has the same client composition as B but enables restricted routing. Group D is an external FedProx baseline with λcream = 0.0 and proximal coefficient μ = 0.01; it changes the local objective and is therefore not a causal, same-loss ablation of B or C.", config)

    add_section_break(doc, config, columns=1)
    write_table(
        doc,
        ["Group", "Setup", "Best Dice", "Final Dice", "Final HD95 (mm)", "Conflict angle (deg)"],
        [
            ["A", "img, FedAvg", "0.8464 ± 0.0090", "0.8328 ± 0.0147", "12.013 ± 0.663", "--"],
            ["B", "img+mm, FedAvg", "0.8687 ± 0.0133", "0.8649 ± 0.0153", "10.416 ± 0.956", "83.855 ± 2.157"],
            ["C", "img+mm, routing", "0.8689 ± 0.0139", "0.8649 ± 0.0158", "10.288 ± 1.138", "83.778 ± 2.714"],
            ["D", "img+mm, FedProx", "0.8907 ± 0.0032", "0.8843 ± 0.0087", "9.128 ± 1.092", "--"],
        ],
        [600, 1600, 1500, 1500, 2050, 2110],
        config,
        "Table 1. Main results across three seeds. Setup abbreviates image-only as img and multimodal as mm. Higher Dice and lower HD95 are better. Group D has a different local objective.",
    )
    add_section_break(doc, config, columns=2)

    add_heading(doc, "4.2. Results and mechanism analysis", config, level=2)
    add_paragraph(doc, "Table 1 shows that adding multimodal clients improves Groups B and C over the image-only baseline A in Best Dice, Final Dice, and Final HD95. With the same client composition, C retains the Dice of B while lowering Final HD95 by 0.128 mm. The mean conflict angles remain below 90° in both heterogeneous protocols.", config)
    add_paragraph(doc, "Round-wise records support restricted routing as a stabilization mechanism rather than a strong optimization intervention: B and C remain below 90° for most rounds, while C has smoother HD95 behavior. The case-level comparison is sample dependent: Cases 0043 and 0044 favor C, Case 0046 is marginal, and Case 0049 favors B.", config)

    add_section_break(doc, config, columns=1)
    write_table(
        doc,
        ["Setting", "Routing", "Rep. update", "Best Dice", "Final Dice", "HD95 (mm)"],
        [
            ["Full C", "yes", "yes", "0.8689 ± 0.0139", "0.8649 ± 0.0158", "10.288 ± 1.138"],
            ["C without rep. update", "yes", "no", "0.8704 ± 0.0140", "0.8661 ± 0.0071", "10.162 ± 0.709"],
            ["λcream", "Rounds", "Final Dice", "HD95 (mm)", "Conflict angle (deg)", ""],
            ["0.02", "30", "0.830277", "12.025466", "77.150981", ""],
            ["0.10", "30", "0.828993", "12.141644", "79.367924", ""],
            ["0.20", "30", "0.830787", "12.190318", "80.383547", ""],
        ],
        [2450, 950, 1250, 1500, 1600, 1610],
        config,
        "Table 2. Component analysis and fixed-30-round ablation. The representation path has no stable separable gain; increasing λcream is not monotonic.",
        subheads={0: "Component analysis", 2: "Distillation-weight ablation in Group C at 30 rounds"},
    )
    add_section_break(doc, config, columns=2)

    add_paragraph(doc, "Table 2 separates routing from the auxiliary representation path. Removing that path does not stably decrease Dice or HD95, so the observed B-C difference is better explained by parameter-mixing constraints. Likewise, the λcream ablation does not improve monotonically; the close B/C Dice values should not be attributed simply to a small distillation weight. Finally, D has higher Dice and lower HD95 than the internal protocols. Since it removes CREAM/alignment loss, this result demonstrates competitive client-drift suppression but cannot isolate FedProx as the only causal factor.", config)
    add_heading(doc, "5. Conclusion", config)
    add_paragraph(doc, "FedSAM3-Hetero makes missing-modality heterogeneous FL explicit at the parameter-group level. Multimodal clients improve the image-only baseline, while restricted routing mainly constrains parameter mixing and offers limited boundary-quality and stability benefits without a substantial Dice gap. The auxiliary representation update is not a stable independent contributor, and the external FedProx outcome does not establish superiority over robust federated optimization. Future work should test the routing contract across additional client compositions, tasks, and multi-center datasets.", config)

    add_heading(doc, "References", config)
    for index, key in enumerate(order, start=1):
        reference = format_reference(index, bibliography[key])
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.first_line_indent = Inches(-0.2)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(reference)
        set_run_font(run, config["body_font"], 7.5)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    return output_docx


def main() -> None:
    arguments = parse_arguments()
    config = read_config(arguments.config)
    output = build_document(config)
    print(output)


if __name__ == "__main__":
    main()
